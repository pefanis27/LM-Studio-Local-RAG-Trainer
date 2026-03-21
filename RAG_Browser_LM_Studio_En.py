from __future__ import annotations

"""
Single-file RAG application for local interaction with LM Studio.

Main features:
- Automatically opens the Streamlit webpage in the default browser
- Selects an LLM from a list and starts the service + loads the model
- Prioritizes background service startup without visibly opening LM Studio
- Scrollable answer panel with proper markdown formatting and a copy button
- Attaches files to the question (PDF, DOCX, code, text, etc.)
- Creates and uses knowledge collections with RAG search
- Cleaner display of sources and collections
"""

import csv
import hashlib
import html
import io
import json
import os
import re
import shutil
import socket
import subprocess
import sys
import threading
import time
import webbrowser
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable
from urllib.parse import urljoin, urlparse


# =========================
# Auto-start via Streamlit
# =========================
STREAMLIT_ENV_FLAG = "LMSTUDIO_RAG_STREAMLIT_ACTIVE"
STREAMLIT_SERVER_PORT = int(os.environ.get("LMSTUDIO_RAG_PORT", "8501"))
STREAMLIT_HOST = os.environ.get("LMSTUDIO_RAG_HOST", "127.0.0.1")
AUTO_OPEN_BROWSER = os.environ.get("LMSTUDIO_RAG_AUTO_OPEN_BROWSER", "1") == "1"
AUTO_START_LM_STUDIO = os.environ.get("LMSTUDIO_RAG_AUTO_START_LMSTUDIO", "1") == "1"
OPEN_LMSTUDIO_APP_IN_BACKGROUND = os.environ.get("LMSTUDIO_RAG_OPEN_APP_IN_BACKGROUND", "1") == "1"
LMSTUDIO_START_TIMEOUT = int(os.environ.get("LMSTUDIO_RAG_LMSTUDIO_START_TIMEOUT", "45"))


def _wait_for_server_and_open_browser(host: str, port: int, timeout: int = 30) -> None:
    deadline = time.time() + timeout
    opened = False
    url = f"http://{host}:{port}"

    while time.time() < deadline and not opened:
        sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        sock.settimeout(0.6)
        try:
            if sock.connect_ex((host, port)) == 0:
                webbrowser.open(url, new=2)
                opened = True
                break
        except Exception:
            pass
        finally:
            try:
                sock.close()
            except Exception:
                pass
        time.sleep(0.5)


if __name__ == "__main__" and os.environ.get(STREAMLIT_ENV_FLAG) != "1":
    env = os.environ.copy()
    env[STREAMLIT_ENV_FLAG] = "1"
    env["LMSTUDIO_RAG_PORT"] = str(STREAMLIT_SERVER_PORT)
    env["LMSTUDIO_RAG_HOST"] = STREAMLIT_HOST

    cmd = [
        sys.executable,
        "-m",
        "streamlit",
        "run",
        os.path.abspath(__file__),
        "--server.headless",
        "true",
        "--server.port",
        str(STREAMLIT_SERVER_PORT),
        "--browser.gatherUsageStats",
        "false",
    ]

    print("🚀 Launching single-file RAG application via Streamlit...")
    print(f"📂 File: {os.path.abspath(__file__)}")
    print(f"🐍 Python: {sys.executable}")
    print(f"🌐 The app will open automatically at: http://{STREAMLIT_HOST}:{STREAMLIT_SERVER_PORT}")
    print()

    if AUTO_OPEN_BROWSER:
        threading.Thread(
            target=_wait_for_server_and_open_browser,
            args=(STREAMLIT_HOST, STREAMLIT_SERVER_PORT, 30),
            daemon=True,
        ).start()

    try:
        raise SystemExit(subprocess.run(cmd, env=env).returncode)
    except FileNotFoundError:
        print("❌ The streamlit command was not found. Install it with:")
        print("   pip install streamlit")
        raise SystemExit(1)
    except KeyboardInterrupt:
        print("\n⛔ Execution was interrupted by the user.")
        raise SystemExit(130)


import numpy as np
import requests
import streamlit as st
from docx import Document
from pypdf import PdfReader

try:
    import markdown as markdown_lib
except Exception:
    markdown_lib = None


# =========================
# Application settings
# =========================
APP_TITLE = "LM Studio Local RAG Trainer"
APP_SUBTITLE = "Local RAG with LM Studio on your own files"
DEFAULT_BASE_URL = "http://localhost:1234"
DEFAULT_API_KEY = "lm-studio"
DEFAULT_CHUNK_SIZE = 900
DEFAULT_CHUNK_OVERLAP = 150
DEFAULT_TOP_K = 4
DEFAULT_TEMPERATURE = 0.2
DEFAULT_MAX_TOKENS = 900
DEFAULT_EMBED_BATCH_SIZE = 16
AUTO_OPTION = "🤖 Automatic selection"
NO_COLLECTION_OPTION = "📎 No collection (attachments only)"
SUPPORTED_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".rst",
    ".log",
    ".pdf",
    ".docx",
    ".rtf",
    ".json",
    ".jsonl",
    ".csv",
    ".tsv",
    ".html",
    ".htm",
    ".xml",
    ".yaml",
    ".yml",
    ".toml",
    ".ini",
    ".cfg",
    ".env",
    ".sql",
    ".tex",
    ".py",
    ".ipynb",
    ".js",
    ".ts",
    ".tsx",
    ".jsx",
    ".java",
    ".c",
    ".cpp",
    ".cc",
    ".h",
    ".hpp",
    ".cs",
    ".go",
    ".rs",
    ".php",
    ".rb",
    ".swift",
    ".kt",
    ".scala",
    ".sh",
    ".bat",
    ".ps1",
}

EMBEDDING_NAME_HINTS = (
    "embedding",
    "embed",
    "embeddinggemma",
    "bge",
    "e5",
    "gte",
    "mxbai",
    "nomic-embed",
    "snowflake-arctic-embed",
    "jina-embeddings",
    "minilm",
    "gte-",
)

BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "data"
COLLECTIONS_DIR = DATA_DIR / "collections"
COLLECTIONS_DIR.mkdir(parents=True, exist_ok=True)


# =========================
# Exceptions
# =========================
class DocumentLoaderError(Exception):
    pass


class LMStudioError(Exception):
    pass


class RAGEngineError(Exception):
    pass


# =========================
# Dataclasses
# =========================
@dataclass(slots=True)
class LoadedDocument:
    filename: str
    text: str


@dataclass(slots=True)
class TextChunk:
    chunk_id: str
    filename: str
    text: str
    chunk_index: int


@dataclass(slots=True)
class SearchResult:
    score: float
    filename: str
    chunk_id: str
    chunk_index: int
    text: str


@dataclass(slots=True)
class AvailableModel:
    identifier: str
    type: str
    display_name: str | None = None
    publisher: str | None = None
    arch: str | None = None
    is_loaded: bool = False


# =========================
# Document helper functions
# =========================
def is_supported(filename: str) -> bool:
    return Path(filename).suffix.lower() in SUPPORTED_EXTENSIONS


def normalize_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[\t\f\v]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" {2,}", " ", text)
    return text.strip()


def _read_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    pages: list[str] = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def _read_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def _read_csv(file_bytes: bytes) -> str:
    text = file_bytes.decode("utf-8", errors="ignore")
    rows = csv.reader(io.StringIO(text))
    rendered_rows = []
    for row in rows:
        rendered_rows.append(" | ".join(cell.strip() for cell in row))
    return "\n".join(rendered_rows)


def _read_json(file_bytes: bytes) -> str:
    obj = json.loads(file_bytes.decode("utf-8", errors="ignore"))
    return json.dumps(obj, ensure_ascii=False, indent=2)


def _read_ipynb(file_bytes: bytes) -> str:
    notebook = json.loads(file_bytes.decode("utf-8", errors="ignore"))
    rendered: list[str] = []
    for idx, cell in enumerate(notebook.get("cells", []), start=1):
        cell_type = cell.get("cell_type", "cell")
        source = "".join(cell.get("source", []))
        if source.strip():
            rendered.append(f"## {cell_type.upper()} CELL {idx}\n{source}")
    metadata = notebook.get("metadata") or {}
    kernel_name = ((metadata.get("kernelspec") or {}).get("display_name") or "").strip()
    prefix = f"Notebook kernel: {kernel_name}\n\n" if kernel_name else ""
    return prefix + "\n\n".join(rendered)


def _read_delimited(file_bytes: bytes, delimiter: str = "\t") -> str:
    text = file_bytes.decode("utf-8", errors="ignore")
    rows = csv.reader(io.StringIO(text), delimiter=delimiter)
    rendered_rows = []
    for row in rows:
        rendered_rows.append(" | ".join(cell.strip() for cell in row))
    return "\n".join(rendered_rows)


def _read_text(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore")


def render_markdown_html(text: str) -> str:
    if markdown_lib is not None:
        try:
            return markdown_lib.markdown(
                text,
                extensions=["extra", "tables", "fenced_code", "sane_lists", "nl2br"],
            )
        except Exception:
            pass

    escaped = html.escape(text)
    return f"<p>{escaped.replace(chr(10), '<br>')}</p>"


def build_attachment_context(
    documents: list[LoadedDocument],
    *,
    start_index: int = 1,
    max_chars_per_file: int = 7000,
    max_total_chars: int = 18000,
) -> tuple[list[str], list[dict]]:
    context_blocks: list[str] = []
    sources: list[dict] = []
    used_total = 0

    for doc_offset, document in enumerate(documents, start=start_index):
        remaining = max_total_chars - used_total
        if remaining <= 0:
            break

        allowed = min(max_chars_per_file, remaining)
        snippet = document.text[:allowed]
        truncated = len(document.text) > len(snippet)
        if truncated:
            snippet = snippet.rstrip() + "\n\n[... the attachment was truncated to fit in the prompt ...]"

        context_blocks.append(
            f"[SOURCE {doc_offset}] file={document.filename} | kind=attachment | score=1.0000\n{snippet}"
        )
        sources.append(
            {
                "score": 1.0,
                "filename": document.filename,
                "chunk_id": f"attachment_{doc_offset:04d}",
                "chunk_index": 0,
                "text": snippet,
                "source_type": "attachment",
            }
        )
        used_total += len(snippet)

    return context_blocks, sources


def load_question_attachments(uploaded_files: Iterable) -> tuple[list[LoadedDocument], list[str]]:
    documents: list[LoadedDocument] = []
    warnings: list[str] = []

    for uploaded in uploaded_files or []:
        if not is_supported(uploaded.name):
            warnings.append(f"The file '{uploaded.name}' was ignored because its type is not supported yet.")
            continue
        try:
            documents.append(extract_text(uploaded.name, uploaded.getvalue()))
        except Exception as exc:
            warnings.append(str(exc))

    return documents, warnings


def extract_text(filename: str, file_bytes: bytes) -> LoadedDocument:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise DocumentLoaderError(f"Unsupported file type: {suffix}")

    try:
        if suffix == ".pdf":
            text = _read_pdf(file_bytes)
        elif suffix == ".docx":
            text = _read_docx(file_bytes)
        elif suffix == ".csv":
            text = _read_csv(file_bytes)
        elif suffix == ".tsv":
            text = _read_delimited(file_bytes, delimiter="	")
        elif suffix in {".json", ".jsonl"}:
            text = _read_json(file_bytes) if suffix == ".json" else _read_text(file_bytes)
        elif suffix == ".ipynb":
            text = _read_ipynb(file_bytes)
        else:
            text = _read_text(file_bytes)
    except Exception as exc:
        raise DocumentLoaderError(f"Failed to read file '{filename}': {exc}") from exc

    text = normalize_text(text)
    if not text.strip():
        raise DocumentLoaderError(f"The file '{filename}' does not contain readable text.")

    return LoadedDocument(filename=filename, text=text)


def split_text(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    if chunk_size <= 0:
        raise ValueError("chunk_size must be a positive number.")
    if chunk_overlap < 0:
        raise ValueError("chunk_overlap cannot be negative.")
    if chunk_overlap >= chunk_size:
        raise ValueError("chunk_overlap must be smaller than chunk_size.")

    if len(text) <= chunk_size:
        return [text]

    chunks: list[str] = []
    start = 0
    text_len = len(text)

    while start < text_len:
        end = min(start + chunk_size, text_len)
        candidate = text[start:end]

        if end < text_len:
            breakpoints = [
                candidate.rfind("\n\n"),
                candidate.rfind("\n"),
                candidate.rfind(". "),
                candidate.rfind(" "),
            ]
            best_break = max(breakpoints)
            if best_break > chunk_size * 0.55:
                end = start + best_break + 1
                candidate = text[start:end]

        chunks.append(candidate.strip())
        if end >= text_len:
            break
        start = max(0, end - chunk_overlap)

    return [chunk for chunk in chunks if chunk]


def build_chunks(documents: Iterable[LoadedDocument], chunk_size: int, chunk_overlap: int) -> list[TextChunk]:
    chunks: list[TextChunk] = []
    for document in documents:
        parts = split_text(document.text, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        for idx, part in enumerate(parts):
            chunks.append(
                TextChunk(
                    chunk_id=f"{Path(document.filename).stem}_{idx:04d}",
                    filename=document.filename,
                    text=part,
                    chunk_index=idx,
                )
            )
    return chunks



# =========================
# LM Studio service helpers
# =========================
def parse_base_url(base_url: str) -> tuple[str, int]:
    raw = base_url.strip() or DEFAULT_BASE_URL
    if "://" not in raw:
        raw = f"http://{raw}"
    parsed = urlparse(raw)
    scheme = parsed.scheme or "http"
    host = parsed.hostname or "127.0.0.1"
    port = parsed.port or (443 if scheme == "https" else 80)
    return host, port


def is_lmstudio_api_reachable(base_url: str, api_key: str | None = None, timeout: int = 4) -> bool:
    headers: dict[str, str] = {}
    if api_key:
        headers["Authorization"] = f"Bearer {api_key}"

    endpoints = ("/api/v1/models", "/v1/models")
    for endpoint in endpoints:
        url = urljoin(base_url.rstrip("/") + "/", endpoint.lstrip("/"))
        try:
            response = requests.get(url, headers=headers, timeout=timeout)
            if response.status_code in {200, 401, 403}:
                return True
        except requests.RequestException:
            continue
    return False


def _spawn_background_process(command: list[str], *, prefer_hidden_app: bool = False) -> bool:
    try:
        kwargs: dict[str, Any] = {
            "stdout": subprocess.DEVNULL,
            "stderr": subprocess.DEVNULL,
        }

        if os.name == "nt":
            creationflags = 0
            creationflags |= getattr(subprocess, "DETACHED_PROCESS", 0)
            creationflags |= getattr(subprocess, "CREATE_NEW_PROCESS_GROUP", 0)

            if not prefer_hidden_app:
                creationflags |= getattr(subprocess, "CREATE_NO_WINDOW", 0)

            startupinfo = subprocess.STARTUPINFO()
            startupinfo.dwFlags |= getattr(subprocess, "STARTF_USESHOWWINDOW", 0)
            if prefer_hidden_app:
                startupinfo.wShowWindow = 0
            else:
                startupinfo.wShowWindow = 0
            kwargs["startupinfo"] = startupinfo
            kwargs["creationflags"] = creationflags
        else:
            kwargs["start_new_session"] = True

        subprocess.Popen(command, **kwargs)
        return True
    except Exception:
        return False


def _spawn_background_app_windows(command: list[str]) -> bool:
    executable = command[0]
    if not Path(executable).exists():
        return False

    if OPEN_LMSTUDIO_APP_IN_BACKGROUND:
        ps_commands = [
            [
                "powershell",
                "-NoProfile",
                "-ExecutionPolicy",
                "Bypass",
                "-WindowStyle",
                "Hidden",
                "-Command",
                f'Start-Process -FilePath "{executable}" -WindowStyle Hidden',
            ],
            [
                "powershell",
                "-NoProfile",
                "-ExecutionPolicy",
                "Bypass",
                "-WindowStyle",
                "Hidden",
                "-Command",
                f'Start-Process -FilePath "{executable}" -WindowStyle Minimized',
            ],
        ]
        for ps_command in ps_commands:
            if _spawn_background_process(ps_command):
                return True

    return _spawn_background_process(command, prefer_hidden_app=True)


def _candidate_lms_commands() -> list[list[str]]:
    candidates: list[list[str]] = []

    which_lms = shutil.which("lms")
    if which_lms:
        candidates.append([which_lms])

    home = Path.home()
    if os.name == "nt":
        candidates.append([str(home / ".lmstudio" / "bin" / "lms.exe")])
    else:
        candidates.append([str(home / ".lmstudio" / "bin" / "lms")])

    seen: set[tuple[str, ...]] = set()
    unique: list[list[str]] = []
    for candidate in candidates:
        key = tuple(candidate)
        if key not in seen:
            seen.add(key)
            unique.append(candidate)
    return unique


def _candidate_lmstudio_app_commands() -> list[list[str]]:
    candidates: list[list[str]] = []

    if os.name == "nt":
        for base in filter(None, [os.environ.get("LOCALAPPDATA"), os.environ.get("PROGRAMFILES"), os.environ.get("PROGRAMFILES(X86)")]):
            exe_path = Path(base) / "Programs" / "LM Studio" / "LM Studio.exe"
            candidates.append([str(exe_path)])

        for base in filter(None, [os.environ.get("PROGRAMFILES"), os.environ.get("PROGRAMFILES(X86)")]):
            exe_path = Path(base) / "LM Studio" / "LM Studio.exe"
            candidates.append([str(exe_path)])
    elif sys.platform == "darwin":
        candidates.extend([
            ["open", "-a", "LM Studio"],
            ["/Applications/LM Studio.app/Contents/MacOS/LM Studio"],
        ])
    else:
        for command_name in ("lm-studio", "LM-Studio", "LM_Studio", "app.lmstudio.LMStudio"):
            exe_path = shutil.which(command_name)
            if exe_path:
                candidates.append([exe_path])

    seen: set[tuple[str, ...]] = set()
    unique: list[list[str]] = []
    for candidate in candidates:
        key = tuple(candidate)
        if key not in seen:
            seen.add(key)
            unique.append(candidate)
    return unique


def _run_command_capture(command: list[str], timeout: int = 20) -> tuple[int, str, str] | None:
    try:
        kwargs: dict[str, Any] = {
            "capture_output": True,
            "text": True,
            "timeout": timeout,
        }
        if os.name == "nt":
            kwargs["creationflags"] = getattr(subprocess, "CREATE_NO_WINDOW", 0)
        completed = subprocess.run(command, **kwargs)
        return completed.returncode, completed.stdout or "", completed.stderr or ""
    except Exception:
        return None


def _unwrap_models_payload(payload: Any) -> list[dict[str, Any]]:
    if isinstance(payload, list):
        return [item for item in payload if isinstance(item, dict)]

    if isinstance(payload, dict):
        collected: list[dict[str, Any]] = []
        for key in ("models", "data", "items", "downloaded_models", "loaded_models", "llm", "embedding", "llms", "embeddings"):
            value = payload.get(key)
            if isinstance(value, list):
                collected.extend(item for item in value if isinstance(item, dict))
        if collected:
            return collected

    return []


def _infer_model_type_for_item(
    identifier: str,
    display_name: str | None = None,
    arch: str | None = None,
    declared_type: str | None = None,
) -> str:
    normalized = (declared_type or "").strip().lower()
    if normalized in {"llm", "embedding"}:
        return normalized

    searchable = " ".join(part for part in [identifier, display_name or "", arch or ""] if part).lower()
    if any(hint in searchable for hint in EMBEDDING_NAME_HINTS):
        return "embedding"
    return "llm"


def _parse_available_model_item(
    item: dict[str, Any],
    *,
    forced_type: str | None = None,
    loaded_identifiers: set[str] | None = None,
) -> AvailableModel | None:
    loaded_identifiers = loaded_identifiers or set()

    loaded_instances = item.get("loaded_instances") or []
    loaded_instance_id = None
    if isinstance(loaded_instances, list) and loaded_instances:
        first_loaded = loaded_instances[0] or {}
        if isinstance(first_loaded, dict):
            loaded_instance_id = (
                first_loaded.get("id")
                or first_loaded.get("instance_id")
                or first_loaded.get("identifier")
            )

    identifier = (
        item.get("key")
        or item.get("id")
        or item.get("identifier")
        or item.get("modelKey")
        or item.get("model_key")
        or item.get("model")
        or item.get("instance_id")
        or loaded_instance_id
        or ""
    )
    identifier = str(identifier).strip()
    if not identifier:
        return None

    display_name = item.get("display_name") or item.get("displayName") or item.get("name") or identifier
    arch = item.get("arch") or item.get("architecture")
    parsed_type = _infer_model_type_for_item(
        identifier=identifier,
        display_name=display_name,
        arch=arch,
        declared_type=forced_type or item.get("type"),
    )

    is_loaded = bool(loaded_instances)
    is_loaded = is_loaded or bool(item.get("loaded")) or bool(item.get("is_loaded"))
    is_loaded = is_loaded or identifier in loaded_identifiers

    return AvailableModel(
        identifier=identifier,
        type=parsed_type,
        display_name=str(display_name),
        publisher=item.get("publisher"),
        arch=str(arch) if arch else None,
        is_loaded=is_loaded,
    )


def list_loaded_model_identifiers_via_cli() -> set[str]:
    identifiers: set[str] = set()

    for base_cmd in _candidate_lms_commands():
        executable = base_cmd[0]
        if not Path(executable).exists() and shutil.which(executable) is None:
            continue

        result = _run_command_capture([executable, "ps", "--json"], timeout=20)
        if not result:
            continue

        returncode, stdout, _ = result
        if returncode != 0 or not stdout.strip():
            continue

        try:
            payload = json.loads(stdout)
        except json.JSONDecodeError:
            continue

        for item in _unwrap_models_payload(payload):
            model = _parse_available_model_item(item)
            if model:
                identifiers.add(model.identifier)
        if identifiers:
            return identifiers

    return identifiers


def list_local_models_via_cli() -> list[AvailableModel]:
    loaded_identifiers = list_loaded_model_identifiers_via_cli()
    aggregated: dict[str, AvailableModel] = {}

    for base_cmd in _candidate_lms_commands():
        executable = base_cmd[0]
        if not Path(executable).exists() and shutil.which(executable) is None:
            continue

        command_specs = [
            ([executable, "ls", "--llm", "--json"], "llm"),
            ([executable, "ls", "--embedding", "--json"], "embedding"),
            ([executable, "ls", "--json"], None),
        ]

        found_any = False
        for command, forced_type in command_specs:
            result = _run_command_capture(command, timeout=30)
            if not result:
                continue

            returncode, stdout, _ = result
            if returncode != 0 or not stdout.strip():
                continue

            try:
                payload = json.loads(stdout)
            except json.JSONDecodeError:
                continue

            items = _unwrap_models_payload(payload)
            if not items:
                continue

            for item in items:
                model = _parse_available_model_item(
                    item,
                    forced_type=forced_type,
                    loaded_identifiers=loaded_identifiers,
                )
                if not model:
                    continue
                existing = aggregated.get(model.identifier)
                if existing is None or (model.is_loaded and not existing.is_loaded):
                    aggregated[model.identifier] = model
            found_any = True

        if found_any:
            break

    return list(aggregated.values())


def launch_lmstudio_app() -> str | None:
    for command in _candidate_lmstudio_app_commands():
        executable = command[0]
        if executable not in {"open"} and not Path(executable).exists():
            continue

        started = False
        if os.name == "nt" and executable.lower().endswith("lm studio.exe"):
            started = _spawn_background_app_windows(command)
        else:
            started = _spawn_background_process(command, prefer_hidden_app=True)

        if started:
            suffix = " in the background" if OPEN_LMSTUDIO_APP_IN_BACKGROUND else ""
            return f"Attempted to open LM Studio{suffix}: {' '.join(command)}"
    return None


def start_lmstudio_server(base_url: str) -> str | None:
    _, port = parse_base_url(base_url)

    for base_cmd in _candidate_lms_commands():
        executable = base_cmd[0]
        if not Path(executable).exists() and shutil.which(executable) is None:
            continue

        command = [executable, "server", "start", "--port", str(port)]
        if _spawn_background_process(command):
            return f"Attempted to start the service: {' '.join(command)}"
    return None


def ensure_lmstudio_online(
    base_url: str,
    api_key: str | None = None,
    wait_timeout: int = LMSTUDIO_START_TIMEOUT,
) -> dict[str, Any]:
    host, port = parse_base_url(base_url)
    report: dict[str, Any] = {
        "ok": False,
        "host": host,
        "port": port,
        "actions": [],
        "message": "",
    }

    if is_lmstudio_api_reachable(base_url, api_key=api_key, timeout=4):
        report["ok"] = True
        report["message"] = "The LM Studio server is already responding."
        report["actions"].append(report["message"])
        return report

    start_message = start_lmstudio_server(base_url)
    if start_message:
        report["actions"].append(start_message)

    early_deadline = time.time() + min(10, max(5, wait_timeout // 3))
    while time.time() < early_deadline:
        if is_lmstudio_api_reachable(base_url, api_key=api_key, timeout=4):
            report["ok"] = True
            report["message"] = f"The LM Studio service started without bringing the application window to the foreground ({host}:{port})."
            report["actions"].append(report["message"])
            return report
        time.sleep(1.0)

    launch_message = launch_lmstudio_app()
    if launch_message:
        report["actions"].append(launch_message)

    if launch_message and not start_message:
        retry_start_message = start_lmstudio_server(base_url)
        if retry_start_message:
            report["actions"].append(retry_start_message)

    deadline = time.time() + max(5, wait_timeout)
    while time.time() < deadline:
        if is_lmstudio_api_reachable(base_url, api_key=api_key, timeout=4):
            report["ok"] = True
            report["message"] = f"LM Studio is available at {host}:{port}."
            report["actions"].append(report["message"])
            return report
        time.sleep(1.0)

    report["message"] = (
        "Could not automatically enable the LM Studio service in the background. "
        "Check whether LM Studio is installed and whether the CLI 'lms' is available."
    )
    if not report["actions"]:
        report["actions"].append("No valid command or executable was found for automatic startup.")
    return report


# =========================
# LM Studio client
# =========================
class LMStudioClient:
    def __init__(self, base_url: str, api_key: str | None = None, timeout: int = 180) -> None:
        self.base_url = base_url.rstrip("/")
        self.api_key = api_key or ""
        self.timeout = timeout

    @property
    def headers(self) -> dict[str, str]:
        headers = {"Content-Type": "application/json"}
        if self.api_key:
            headers["Authorization"] = f"Bearer {self.api_key}"
        return headers

    def _infer_model_type(self, identifier: str, display_name: str | None, arch: str | None, declared_type: str | None) -> str:
        return _infer_model_type_for_item(identifier, display_name, arch, declared_type)

    def _get(self, path: str) -> Any:
        url = urljoin(self.base_url + "/", path.lstrip("/"))
        try:
            response = requests.get(url, headers=self.headers, timeout=self.timeout)
            response.raise_for_status()
            return response.json()
        except requests.RequestException as exc:
            raise LMStudioError(f"GET request to LM Studio failed: {exc}") from exc

    def _post(self, path: str, payload: dict[str, Any]) -> Any:
        url = urljoin(self.base_url + "/", path.lstrip("/"))
        response = None
        try:
            response = requests.post(url, json=payload, headers=self.headers, timeout=self.timeout)
            response.raise_for_status()
            if response.text.strip():
                return response.json()
            return {}
        except requests.RequestException as exc:
            details = ""
            try:
                if response is not None:
                    details = f" | Response: {response.text[:500]}"
            except Exception:
                pass
            raise LMStudioError(f"POST request to LM Studio failed: {exc}{details}") from exc

    def list_models(self) -> list[AvailableModel]:
        payload = self._get("/api/v1/models")
        items = _unwrap_models_payload(payload)

        unique: dict[str, AvailableModel] = {}
        for item in items:
            model = _parse_available_model_item(item)
            if not model:
                continue
            existing = unique.get(model.identifier)
            if existing is None or (model.is_loaded and not existing.is_loaded):
                unique[model.identifier] = model

        return list(unique.values())

    def load_model(self, model: str) -> dict[str, Any]:
        return self._post("/api/v1/models/load", {"model": model})

    def create_embeddings(self, texts: list[str], model: str) -> list[list[float]]:
        if not texts:
            return []
        payload = {
            "model": model,
            "input": [text.replace("\n", " ") for text in texts],
        }
        data = self._post("/v1/embeddings", payload)
        vectors = data.get("data", [])
        if not vectors:
            raise LMStudioError("LM Studio did not return embeddings.")
        return [item["embedding"] for item in vectors]

    def chat(
        self,
        model: str,
        messages: list[dict[str, str]],
        temperature: float = 0.2,
        max_tokens: int = 900,
    ) -> str:
        payload = {
            "model": model,
            "messages": messages,
            "temperature": temperature,
            "max_tokens": max_tokens,
            "stream": False,
        }
        data = self._post("/v1/chat/completions", payload)
        choices = data.get("choices", [])
        if not choices:
            raise LMStudioError("LM Studio did not return a response.")
        return choices[0].get("message", {}).get("content", "").strip()


# =========================
# RAG Engine
# =========================
class RAGEngine:
    def __init__(self, collections_dir: str | Path = COLLECTIONS_DIR) -> None:
        self.collections_dir = Path(collections_dir)
        self.collections_dir.mkdir(parents=True, exist_ok=True)

    def list_collections(self) -> list[str]:
        return sorted([item.name for item in self.collections_dir.iterdir() if item.is_dir()])

    def delete_collection(self, collection_name: str) -> None:
        collection_dir = self._collection_dir(collection_name)
        if collection_dir.exists():
            shutil.rmtree(collection_dir)

    def _collection_dir(self, collection_name: str) -> Path:
        safe_name = collection_name.strip().replace(" ", "_")
        if not safe_name:
            raise RAGEngineError("The collection name cannot be empty.")
        return self.collections_dir / safe_name

    def build_collection(
        self,
        collection_name: str,
        uploaded_files: Iterable,
        client: LMStudioClient,
        embedding_model: str,
        chunk_size: int,
        chunk_overlap: int,
        batch_size: int,
    ) -> dict:
        collection_dir = self._collection_dir(collection_name)
        source_dir = collection_dir / "source_files"
        source_dir.mkdir(parents=True, exist_ok=True)

        documents: list[LoadedDocument] = []
        saved_files: list[str] = []

        for uploaded in uploaded_files:
            file_bytes = uploaded.getvalue()
            (source_dir / uploaded.name).write_bytes(file_bytes)
            loaded = extract_text(uploaded.name, file_bytes)
            documents.append(loaded)
            saved_files.append(uploaded.name)

        if not documents:
            raise RAGEngineError("No files were provided to create the knowledge base.")

        chunks = build_chunks(documents, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        if not chunks:
            raise RAGEngineError("No chunks were produced from your data.")

        chunk_texts = [chunk.text for chunk in chunks]
        all_embeddings: list[list[float]] = []

        for start in range(0, len(chunk_texts), batch_size):
            batch = chunk_texts[start : start + batch_size]
            batch_vectors = client.create_embeddings(batch, model=embedding_model)
            all_embeddings.extend(batch_vectors)

        matrix = np.array(all_embeddings, dtype=np.float32)
        normalized = self._normalize(matrix)

        metadata = {
            "collection_name": collection_dir.name,
            "created_at": datetime.now().isoformat(timespec="seconds"),
            "embedding_model": embedding_model,
            "chunk_size": chunk_size,
            "chunk_overlap": chunk_overlap,
            "documents": saved_files,
            "num_chunks": len(chunks),
            "embedding_dim": int(normalized.shape[1]),
        }

        chunks_payload = [
            {
                "chunk_id": chunk.chunk_id,
                "filename": chunk.filename,
                "chunk_index": chunk.chunk_index,
                "text": chunk.text,
            }
            for chunk in chunks
        ]

        np.save(collection_dir / "embeddings.npy", normalized)
        (collection_dir / "chunks.json").write_text(json.dumps(chunks_payload, ensure_ascii=False, indent=2), encoding="utf-8")
        (collection_dir / "metadata.json").write_text(json.dumps(metadata, ensure_ascii=False, indent=2), encoding="utf-8")

        return metadata

    def get_collection_metadata(self, collection_name: str) -> dict:
        collection_dir = self._collection_dir(collection_name)
        metadata_path = collection_dir / "metadata.json"
        if not metadata_path.exists():
            raise RAGEngineError(f"No metadata was found for collection '{collection_name}'.")
        return json.loads(metadata_path.read_text(encoding="utf-8"))

    def search(
        self,
        collection_name: str,
        query: str,
        client: LMStudioClient,
        embedding_model: str,
        top_k: int,
    ) -> list[SearchResult]:
        collection_dir = self._collection_dir(collection_name)
        chunks_path = collection_dir / "chunks.json"
        vectors_path = collection_dir / "embeddings.npy"

        if not chunks_path.exists() or not vectors_path.exists():
            raise RAGEngineError(f"Collection '{collection_name}' is not ready for search.")

        chunk_items = json.loads(chunks_path.read_text(encoding="utf-8"))
        vectors = np.load(vectors_path)

        query_vector = np.array(client.create_embeddings([query], model=embedding_model)[0], dtype=np.float32)
        query_vector = self._normalize(query_vector.reshape(1, -1))[0]

        scores = vectors @ query_vector
        if scores.ndim != 1:
            scores = scores.reshape(-1)

        top_indices = np.argsort(scores)[::-1][: max(1, top_k)]

        results: list[SearchResult] = []
        for idx in top_indices:
            item = chunk_items[int(idx)]
            results.append(
                SearchResult(
                    score=float(scores[int(idx)]),
                    filename=item["filename"],
                    chunk_id=item["chunk_id"],
                    chunk_index=int(item["chunk_index"]),
                    text=item["text"],
                )
            )
        return results

    def answer_question(
        self,
        collection_name: str | None,
        question: str,
        client: LMStudioClient,
        llm_model: str,
        embedding_model: str | None,
        top_k: int,
        temperature: float,
        max_tokens: int,
        attachment_documents: list[LoadedDocument] | None = None,
    ) -> dict:
        hits: list[SearchResult] = []
        context_blocks: list[str] = []
        sources: list[dict] = []

        if collection_name:
            if not embedding_model:
                raise RAGEngineError("An embedding model is required when a knowledge collection is used.")

            hits = self.search(
                collection_name=collection_name,
                query=question,
                client=client,
                embedding_model=embedding_model,
                top_k=top_k,
            )

            for index, hit in enumerate(hits, start=1):
                context_blocks.append(
                    f"[SOURCE {index}] file={hit.filename} | chunk={hit.chunk_index} | score={hit.score:.4f}\n{hit.text}"
                )
            sources.extend([{**asdict(hit), "source_type": "collection"} for hit in hits])

        attachment_documents = attachment_documents or []
        if attachment_documents:
            attachment_blocks, attachment_sources = build_attachment_context(
                attachment_documents,
                start_index=len(context_blocks) + 1,
            )
            context_blocks.extend(attachment_blocks)
            sources.extend(attachment_sources)

        if not context_blocks:
            raise RAGEngineError(
                "No relevant excerpts were found and no readable attachments were provided to answer the question."
            )

        system_prompt = (
            "You are a RAG assistant that answers exclusively based on the provided context. "
            "If the context is insufficient, say so clearly. "
            "Answer in English, be clear, and do not invent information. "
            "When you use details from the context, cite them in the form [SOURCE 1], [SOURCE 2]."
        )

        joined_context = "\n\n".join(context_blocks)
        user_prompt = (
            f"Context:\n\n{joined_context}\n\n"
            f"User question:\n{question}\n\n"
            "Instructions:\n"
            "1. Answer only with information supported by the context.\n"
            "2. If information is missing, state that clearly.\n"
            "3. Include citations [SOURCE x] inside the answer.\n"
            "4. Use bullets or short sections when they improve clarity.\n"
            "5. If there are attachments, you may use them as additional context sources."
        )

        answer = client.chat(
            model=llm_model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=temperature,
            max_tokens=max_tokens,
        )

        return {
            "answer": answer,
            "sources": sources,
        }

    @staticmethod
    def _normalize(vectors: np.ndarray) -> np.ndarray:
        norms = np.linalg.norm(vectors, axis=1, keepdims=True)
        norms[norms == 0.0] = 1.0
        return vectors / norms


# =========================
# UI helpers
# =========================
@st.cache_resource

def get_engine() -> RAGEngine:
    return RAGEngine()



def inject_custom_css() -> None:
    st.markdown(
        """
        <style>
        :root {
            --app-primary: #4f46e5;
            --app-secondary: #0ea5e9;
            --app-accent: #14b8a6;
            --app-border: rgba(148, 163, 184, 0.20);
            --app-glass: rgba(15, 23, 42, 0.035);
            --app-text: #0f172a;
            --app-muted: #475569;
        }

        .stApp {
            background:
                radial-gradient(circle at top left, rgba(79, 70, 229, 0.18), transparent 28%),
                radial-gradient(circle at top right, rgba(14, 165, 233, 0.16), transparent 22%),
                linear-gradient(180deg, rgba(248,250,252,1) 0%, rgba(241,245,249,1) 100%);
        }

        [data-testid="stSidebar"] {
            background: linear-gradient(180deg, rgba(15,23,42,0.96), rgba(30,41,59,0.96));
            border-right: 1px solid rgba(255,255,255,0.08);
        }

        [data-testid="stSidebar"] [data-testid="stMarkdownContainer"],
        [data-testid="stSidebar"] label,
        [data-testid="stSidebar"] .stCaption,
        [data-testid="stSidebar"] .stText,
        [data-testid="stSidebar"] p,
        [data-testid="stSidebar"] h1,
        [data-testid="stSidebar"] h2,
        [data-testid="stSidebar"] h3 {
            color: #e5eefc !important;
        }

        .hero-card {
            padding: 1.05rem 1.15rem;
            border-radius: 24px;
            border: 1px solid rgba(255,255,255,0.22);
            background: linear-gradient(135deg, rgba(79,70,229,0.96), rgba(14,165,233,0.92));
            color: white;
            box-shadow: 0 22px 50px rgba(37, 99, 235, 0.18);
            margin-bottom: 1rem;
        }

        .hero-title {
            font-size: 1.55rem;
            font-weight: 800;
            line-height: 1.1;
            margin-bottom: 0.35rem;
        }

        .hero-subtitle {
            font-size: 0.93rem;
            opacity: 0.95;
            line-height: 1.45;
        }

        .mini-card {
            background: rgba(255,255,255,0.72);
            backdrop-filter: blur(10px);
            border: 1px solid var(--app-border);
            border-radius: 20px;
            padding: 0.82rem 0.88rem 0.78rem 0.88rem;
            box-shadow: 0 8px 25px rgba(15, 23, 42, 0.06);
            margin-bottom: 0.9rem;
        }

        .mini-card-title {
            font-size: 0.78rem;
            color: var(--app-muted);
            margin-bottom: 0.25rem;
            word-break: break-word;
        }

        .mini-card-value {
            font-size: 1.02rem;
            font-weight: 800;
            color: var(--app-text);
        }

        .stTextInput input,
        .stTextArea textarea,
        .stSelectbox [data-baseweb="select"],
        .stNumberInput input,
        .stFileUploader,
        [data-testid="stVerticalBlockBorderWrapper"],
        [data-testid="stMetric"],
        .stAlert {
            border: 1px solid rgba(148, 163, 184, 0.18) !important;
            border-radius: 14px !important;
        }

        .section-note {
            padding: 0.95rem 1rem;
            border-radius: 18px;
            border: 1px solid rgba(59,130,246,0.18);
            background: rgba(255,255,255,0.75);
            margin-bottom: 0.9rem;
        }

        .answer-shell {
            border-radius: 24px;
            border: 1px solid rgba(79, 70, 229, 0.14);
            background: linear-gradient(180deg, rgba(255,255,255,0.92), rgba(248,250,252,0.92));
            padding: 0.85rem;
            box-shadow: 0 10px 24px rgba(15, 23, 42, 0.06);
        }

        .answer-header {
            display: flex;
            align-items: center;
            justify-content: space-between;
            gap: 1rem;
            margin-bottom: 0.65rem;
        }

        .answer-title {
            font-size: 0.96rem;
            font-weight: 800;
            color: #111827;
        }

        .answer-badge {
            padding: 0.28rem 0.7rem;
            border-radius: 999px;
            background: rgba(79,70,229,0.1);
            color: #4338ca;
            font-size: 0.76rem;
            font-weight: 700;
        }

        .status-pill {
            display: inline-block;
            padding: 0.34rem 0.7rem;
            border-radius: 999px;
            background: rgba(15,23,42,0.06);
            border: 1px solid rgba(148,163,184,0.25);
            color: var(--app-text);
            margin-right: 0.45rem;
            margin-bottom: 0.45rem;
            font-size: 0.78rem;
            font-weight: 600;
        }

        .stButton > button,
        div[data-testid="stDownloadButton"] > button,
        div[data-testid="stFormSubmitButton"] > button {
            width: 100%;
            min-height: 2.75rem;
            border-radius: 16px;
            border: 1px solid rgba(79, 70, 229, 0.14);
            background: linear-gradient(180deg, rgba(255,255,255,0.98), rgba(238,242,255,0.98));
            color: #0f172a !important;
            -webkit-text-fill-color: #0f172a !important;
            font-weight: 800;
            font-size: 0.86rem;
            box-shadow: 0 10px 24px rgba(15, 23, 42, 0.08);
            transition: all 0.18s ease;
            opacity: 1 !important;
        }

        .stButton > button *,
        div[data-testid="stDownloadButton"] > button *,
        div[data-testid="stFormSubmitButton"] > button * {
            color: #0f172a !important;
            -webkit-text-fill-color: #0f172a !important;
            opacity: 1 !important;
        }

        .stButton > button p,
        .stButton > button span,
        .stButton > button div,
        div[data-testid="stDownloadButton"] > button p,
        div[data-testid="stDownloadButton"] > button span,
        div[data-testid="stDownloadButton"] > button div,
        div[data-testid="stFormSubmitButton"] > button p,
        div[data-testid="stFormSubmitButton"] > button span,
        div[data-testid="stFormSubmitButton"] > button div {
            color: #0f172a !important;
            -webkit-text-fill-color: #0f172a !important;
            margin: 0 !important;
            line-height: 1.25 !important;
        }

        .stButton > button:hover,
        div[data-testid="stDownloadButton"] > button:hover,
        div[data-testid="stFormSubmitButton"] > button:hover {
            transform: translateY(-1px);
            border-color: rgba(79, 70, 229, 0.38);
            box-shadow: 0 14px 28px rgba(79, 70, 229, 0.12);
        }

        .stButton > button[kind="primary"],
        div[data-testid="stFormSubmitButton"] > button[kind="primary"] {
            background: linear-gradient(135deg, rgba(79,70,229,1), rgba(14,165,233,1));
            color: #ffffff !important;
            -webkit-text-fill-color: #ffffff !important;
            border: none;
        }

        .stButton > button[kind="primary"] *,
        div[data-testid="stFormSubmitButton"] > button[kind="primary"] * {
            color: #ffffff !important;
            -webkit-text-fill-color: #ffffff !important;
        }

        [data-testid="stSidebar"] .stButton > button,
        [data-testid="stSidebar"] div[data-testid="stDownloadButton"] > button,
        [data-testid="stSidebar"] div[data-testid="stFormSubmitButton"] > button {
            background: linear-gradient(180deg, rgba(255,255,255,0.98), rgba(226,232,240,0.98));
            color: #0f172a !important;
            -webkit-text-fill-color: #0f172a !important;
            border: 1px solid rgba(148,163,184,0.28);
        }

        [data-testid="stSidebar"] .stButton > button *,
        [data-testid="stSidebar"] div[data-testid="stDownloadButton"] > button *,
        [data-testid="stSidebar"] div[data-testid="stFormSubmitButton"] > button * {
            color: #0f172a !important;
            -webkit-text-fill-color: #0f172a !important;
        }

        [data-testid="stSidebar"] .stButton > button[kind="primary"],
        [data-testid="stSidebar"] div[data-testid="stFormSubmitButton"] > button[kind="primary"] {
            background: linear-gradient(135deg, rgba(79,70,229,1), rgba(14,165,233,1));
            color: #ffffff !important;
            -webkit-text-fill-color: #ffffff !important;
            border: none;
        }

        [data-testid="stSidebar"] .stButton > button[kind="primary"] *,
        [data-testid="stSidebar"] div[data-testid="stFormSubmitButton"] > button[kind="primary"] * {
            color: #ffffff !important;
            -webkit-text-fill-color: #ffffff !important;
        }

        div[data-testid="stTabs"] button {
            font-weight: 700;
            font-size: 0.90rem;
            color: #0f172a !important;
            -webkit-text-fill-color: #0f172a !important;
        }

        button[kind="secondary"],
        button[kind="primary"],
        button[data-testid="baseButton-secondary"],
        button[data-testid="baseButton-primary"] {
            opacity: 1 !important;
        }

        button[kind="secondary"] *,
        button[kind="primary"] *,
        button[data-testid="baseButton-secondary"] *,
        button[data-testid="baseButton-primary"] * {
            opacity: 1 !important;
        }

        div[data-testid="stCodeBlock"] {
            border-radius: 18px !important;
            border: 1px solid rgba(148,163,184,0.2);
        }

        .rag-answer-inline {
            border: 1px solid rgba(148,163,184,0.22);
            border-radius: 16px;
            background: rgba(255,255,255,0.92);
            padding: 0.9rem 0.95rem;
        }

        .rag-answer-inline-toolbar {
            display: flex;
            align-items: center;
            justify-content: space-between;
            gap: 0.75rem;
            margin-bottom: 0.65rem;
            flex-wrap: wrap;
        }

        .rag-answer-inline-note {
            font-size: 0.80rem;
            color: #475569;
            font-weight: 700;
        }

        .rag-answer-scroll-inline {
            height: 360px;
            overflow-y: auto;
            overflow-x: auto;
            border: 1px solid rgba(148,163,184,0.24);
            border-radius: 14px;
            background: rgba(255,255,255,0.98);
            padding: 0.95rem 1rem 1.1rem 1rem;
            line-height: 1.62;
            word-break: break-word;
        }

        .rag-answer-scroll-inline h1,
        .rag-answer-scroll-inline h2,
        .rag-answer-scroll-inline h3,
        .rag-answer-scroll-inline h4,
        .rag-answer-scroll-inline h5,
        .rag-answer-scroll-inline h6 {
            color: #111827;
            margin-top: 0.75rem;
            margin-bottom: 0.35rem;
            line-height: 1.28;
        }

        .rag-answer-scroll-inline p,
        .rag-answer-scroll-inline li {
            color: #0f172a;
            font-size: 0.93rem;
        }

        .rag-answer-scroll-inline ul,
        .rag-answer-scroll-inline ol {
            padding-left: 1.15rem;
        }

        .rag-answer-scroll-inline code {
            font-family: Consolas, Menlo, monospace;
            background: rgba(15,23,42,0.06);
            padding: 0.10rem 0.3rem;
            border-radius: 6px;
            font-size: 0.90em;
        }

        .rag-answer-scroll-inline pre {
            white-space: pre-wrap;
            word-break: break-word;
            background: #0f172a;
            color: #e2e8f0;
            padding: 0.85rem;
            border-radius: 12px;
            overflow-x: auto;
        }

        .rag-answer-scroll-inline pre code {
            background: transparent;
            color: inherit;
            padding: 0;
        }

        .rag-answer-scroll-inline blockquote {
            margin: 0.65rem 0;
            padding: 0.1rem 0 0.1rem 0.8rem;
            border-left: 4px solid rgba(79,70,229,0.35);
            color: #334155;
            background: rgba(79,70,229,0.04);
        }

        .rag-answer-scroll-inline table {
            width: 100%;
            border-collapse: collapse;
            margin: 0.75rem 0;
        }

        .rag-answer-scroll-inline th,
        .rag-answer-scroll-inline td {
            border: 1px solid rgba(148,163,184,0.32);
            padding: 0.48rem 0.58rem;
            text-align: left;
            vertical-align: top;
        }

        .rag-answer-scroll-inline th {
            background: rgba(15,23,42,0.04);
        }

        .footer-note {
            border-radius: 18px;
            padding: 1rem 1.1rem;
            background: rgba(255,255,255,0.72);
            border: 1px solid var(--app-border);
        }

        .service-box {
            border-radius: 18px;
            padding: 0.9rem 1rem;
            background: rgba(255,255,255,0.08);
            border: 1px solid rgba(255,255,255,0.10);
            margin-top: 0.8rem;
            margin-bottom: 0.5rem;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )



def build_client() -> LMStudioClient:
    return LMStudioClient(
        base_url=st.session_state.base_url,
        api_key=st.session_state.api_key,
        timeout=st.session_state.timeout_seconds,
    )



def refresh_models() -> None:
    models: list[AvailableModel] = []
    errors: list[str] = []
    source = "API"

    try:
        client = build_client()
        models = client.list_models()
        if not models:
            errors.append("The API responded but did not return any models.")
    except Exception as exc:
        errors.append(f"API: {exc}")

    if not models:
        try:
            models = list_local_models_via_cli()
            source = "CLI"
        except Exception as exc:
            errors.append(f"CLI: {exc}")

    models = sorted(
        models,
        key=lambda m: (
            0 if m.type == "llm" else 1,
            0 if m.is_loaded else 1,
            (m.display_name or m.identifier).lower(),
        ),
    )

    st.session_state.models_cache = {
        "all": [asdict(model) for model in models],
        "llm": [model.identifier for model in models if model.type == "llm"],
        "embedding": [model.identifier for model in models if model.type == "embedding"],
    }
    st.session_state.models_source = source

    if not models and errors:
        raise LMStudioError(" | ".join(errors))



def get_cached_models(model_type: str | None = None) -> list[AvailableModel]:
    models = [AvailableModel(**item) for item in st.session_state.models_cache.get("all", [])]
    if model_type is None:
        return models
    return [model for model in models if model.type == model_type]



def get_model_by_identifier(identifier: str) -> AvailableModel | None:
    for model in get_cached_models():
        if model.identifier == identifier:
            return model
    return None



def format_model_option(identifier: str) -> str:
    if identifier == AUTO_OPTION:
        return AUTO_OPTION

    model = get_model_by_identifier(identifier)
    if not model:
        return identifier

    parts = [model.display_name or model.identifier]
    meta: list[str] = [model.type]
    if model.is_loaded:
        meta.append("loaded")
    if model.arch:
        meta.append(model.arch)
    return f"{parts[0]} [{model.identifier}] — " + " • ".join(meta)



def ellipsize_middle(value: str, max_length: int = 28) -> str:
    value = str(value or "")
    if len(value) <= max_length:
        return value
    keep = max(6, (max_length - 3) // 2)
    tail = max(6, max_length - keep - 3)
    return f"{value[:keep]}...{value[-tail:]}"



def format_collection_option(option: str) -> str:
    if option == NO_COLLECTION_OPTION:
        return option
    return ellipsize_middle(option, max_length=26)



def show_sources(sources: list[dict]) -> None:
    for idx, src in enumerate(sources, start=1):
        source_type = src.get("source_type", "collection")
        source_label = "Attachment" if source_type == "attachment" else "Collection"
        chunk_index = src.get("chunk_index", 0)
        score = float(src.get("score", 0.0))
        with st.expander(
            f"SOURCE {idx} — {src['filename']} — {source_label} — chunk {chunk_index} — score {score:.4f}"
        ):
            st.code(src.get("text", ""), language=None)



def rank_embedding_candidate(model: AvailableModel) -> tuple[int, int, str]:
    searchable = " ".join(
        part for part in [model.identifier, model.display_name or "", model.arch or ""] if part
    ).lower()

    for idx, hint in enumerate(EMBEDDING_NAME_HINTS):
        if hint in searchable:
            return (0, idx, searchable)

    return (1, 999, searchable)



def resolve_model(
    *,
    model_type: str,
    selected_identifier: str | None,
    engine: RAGEngine,
    collection_name: str | None = None,
) -> tuple[AvailableModel, str]:
    available = get_cached_models(model_type)
    if not available:
        raise LMStudioError(
            f"No available model of type '{model_type}' was found in LM Studio. "
            "Refresh or download/add the appropriate model."
        )

    if selected_identifier and selected_identifier != AUTO_OPTION:
        model = get_model_by_identifier(selected_identifier)
        if model and model.type == model_type:
            return model, "The model you selected manually is being used."
        raise LMStudioError(f"The selected model '{selected_identifier}' is not available as a {model_type} model.")

    if model_type == "embedding" and collection_name:
        try:
            metadata = engine.get_collection_metadata(collection_name)
            embedded_identifier = metadata.get("embedding_model")
            model = get_model_by_identifier(embedded_identifier) if embedded_identifier else None
            if model and model.type == "embedding":
                return model, "The embedding model used for the collection was found and preferred."
        except Exception:
            pass

    loaded = [model for model in available if model.is_loaded]
    if loaded:
        if model_type == "embedding":
            loaded = sorted(loaded, key=rank_embedding_candidate)
        else:
            loaded = sorted(loaded, key=lambda m: (m.display_name or m.identifier).lower())
        return loaded[0], "A suitable model is already loaded in LM Studio."

    if model_type == "embedding":
        sorted_available = sorted(available, key=rank_embedding_candidate)
        return sorted_available[0], "No manual selection was provided; an automatic fallback embedding model was used."

    sorted_available = sorted(available, key=lambda m: (m.display_name or m.identifier).lower())
    return sorted_available[0], "No manual selection was provided; an automatic fallback LLM was used."



def ensure_model_loaded(client: LMStudioClient, model: AvailableModel) -> str:
    startup_report = ensure_lmstudio_online(
        client.base_url,
        api_key=client.api_key,
        wait_timeout=min(20, LMSTUDIO_START_TIMEOUT),
    )
    st.session_state.lmstudio_boot_report = startup_report
    if not startup_report.get("ok"):
        raise LMStudioError(startup_report.get("message", "The LM Studio service is not available."))

    try:
        refresh_models()
    except Exception:
        pass

    refreshed_before_load = get_model_by_identifier(model.identifier)
    if refreshed_before_load and refreshed_before_load.is_loaded:
        return f"Model '{model.identifier}' is already loaded."

    client.load_model(model.identifier)
    time.sleep(0.8)

    try:
        refresh_models()
    except Exception:
        pass

    refreshed = get_model_by_identifier(model.identifier)
    if refreshed and refreshed.is_loaded:
        return f"Model '{model.identifier}' loaded successfully."

    return f"A load request was sent for model '{model.identifier}'."



def render_hero(collections_count: int, llm_count: int, embedding_count: int) -> None:
    st.markdown(
        f"""
        <div class="hero-card">
            <div class="hero-title">🧠 {APP_TITLE}</div>
            <div class="hero-subtitle">
                Build local knowledge from your files and get RAG answers from LM Studio
                through a cleaner web interface.
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    a, b, c = st.columns(3)
    with a:
        st.markdown(
            f"<div class='mini-card'><div class='mini-card-title'>Knowledge collections</div><div class='mini-card-value'>{collections_count}</div></div>",
            unsafe_allow_html=True,
        )
    with b:
        st.markdown(
            f"<div class='mini-card'><div class='mini-card-title'>LLM models</div><div class='mini-card-value'>{llm_count}</div></div>",
            unsafe_allow_html=True,
        )
    with c:
        st.markdown(
            f"<div class='mini-card'><div class='mini-card-title'>Embedding models</div><div class='mini-card-value'>{embedding_count}</div></div>",
            unsafe_allow_html=True,
        )



def render_status_strip(base_url: str, collections_count: int) -> None:
    st.markdown(
        f"""
        <div class="section-note">
            <span class="status-pill">Base URL: {base_url}</span>
            <span class="status-pill">Port app: {STREAMLIT_SERVER_PORT}</span>
            <span class="status-pill">Collections: {collections_count}</span>
            <span class="status-pill">Auto browser: {'ON' if AUTO_OPEN_BROWSER else 'OFF'}</span>
        </div>
        """,
        unsafe_allow_html=True,
    )



def build_answer_panel_html(answer_text: str, dom_id: str) -> str:
    rendered_html = render_markdown_html(answer_text or "")
    return f"""
    <div class="rag-answer-inline" id="answer_wrap_{dom_id}">
        <div class="rag-answer-inline-toolbar">
            <div class="rag-answer-inline-note">Scrollable answer panel with markdown formatting</div>
        </div>
        <div class="rag-answer-scroll-inline">{rendered_html}</div>
    </div>
    """


def render_answer_panel(result: dict | None, panel_key_prefix: str = "latest") -> None:
    if not result:
        st.info("There is no answer yet. Ask a question from the 'Ask' tab.")
        return

    answer_text = result.get("answer", "")
    question = result.get("question", "")
    metadata = result.get("meta", {})
    result_seq = int(result.get("result_seq", 0) or 0)
    answer_hash = hashlib.md5(f"{answer_text or ''}_{question}_{result_seq}".encode("utf-8")).hexdigest()[:12]
    dom_id = f"{panel_key_prefix}_{answer_hash}_{result_seq}".replace("-", "_")

    st.markdown(
        """
        <div class="answer-shell">
            <div class="answer-header">
                <div class="answer-title">📋 Latest answer</div>
                <div class="answer-badge">Formatted view</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    if question:
        st.caption(f"Question: {question}")

    info_cols = st.columns(5)
    info_cols[0].metric("Collection", metadata.get("collection", "-"))
    info_cols[1].metric("LLM", metadata.get("llm_model", "-"))
    info_cols[2].metric("Embedding", metadata.get("embedding_model", "-"))
    info_cols[3].metric("Sources", str(len(result.get("sources", []))))
    info_cols[4].metric("Attachments", str(metadata.get("attachments_count", 0)))

    st.markdown(build_answer_panel_html(answer_text, dom_id), unsafe_allow_html=True)

    attachment_names = metadata.get("attachments") or []
    if attachment_names:
        st.caption("Question attachments: " + ", ".join(attachment_names))

    with st.expander("Text for copying", expanded=False):
        st.text_area(
            "Plain answer text",
            value=answer_text,
            height=140,
            key=f"copy_answer_text_{panel_key_prefix}_{answer_hash}",
            label_visibility="collapsed",
        )

    btn_col1, btn_col2 = st.columns([1, 1])
    with btn_col1:
        st.download_button(
            "⬇️ Download answer .txt",
            data=answer_text.encode("utf-8"),
            file_name="rag_answer.txt",
            mime="text/plain",
            use_container_width=True,
            key=f"download_answer_{panel_key_prefix}_{answer_hash}",
        )
    with btn_col2:
        if st.button(
            "🧹 Clear answer",
            use_container_width=True,
            key=f"clear_answer_{panel_key_prefix}_{answer_hash}",
        ):
            st.session_state.latest_result = None
            st.rerun()

    st.markdown("### Sources used")
    show_sources(result.get("sources", []))


# =========================
# Streamlit app
# =========================
st.set_page_config(page_title=APP_TITLE, page_icon="🧠", layout="wide")
inject_custom_css()

if "models_cache" not in st.session_state:
    st.session_state.models_cache = {"all": [], "llm": [], "embedding": []}
if "base_url" not in st.session_state:
    st.session_state.base_url = DEFAULT_BASE_URL
if "api_key" not in st.session_state:
    st.session_state.api_key = DEFAULT_API_KEY
if "timeout_seconds" not in st.session_state:
    st.session_state.timeout_seconds = 180
if "startup_refresh_done" not in st.session_state:
    st.session_state.startup_refresh_done = False
if "latest_result" not in st.session_state:
    st.session_state.latest_result = None
if "latest_result_seq" not in st.session_state:
    st.session_state.latest_result_seq = 0
if "lmstudio_boot_report" not in st.session_state:
    st.session_state.lmstudio_boot_report = None
if "lmstudio_auto_boot_done" not in st.session_state:
    st.session_state.lmstudio_auto_boot_done = False
if "models_source" not in st.session_state:
    st.session_state.models_source = "API"
if "service_llm_model" not in st.session_state:
    st.session_state.service_llm_model = AUTO_OPTION
if "ask_llm_model" not in st.session_state:
    st.session_state.ask_llm_model = AUTO_OPTION

engine = get_engine()

if not st.session_state.startup_refresh_done:
    if AUTO_START_LM_STUDIO and not st.session_state.lmstudio_auto_boot_done:
        st.session_state.lmstudio_boot_report = ensure_lmstudio_online(
            st.session_state.base_url,
            api_key=st.session_state.api_key,
            wait_timeout=LMSTUDIO_START_TIMEOUT,
        )
        st.session_state.lmstudio_auto_boot_done = True
    try:
        refresh_models()
    except Exception:
        pass
    st.session_state.startup_refresh_done = True

collections = engine.list_collections()
llm_models = st.session_state.models_cache["llm"]
embedding_models = st.session_state.models_cache["embedding"]

render_hero(
    collections_count=len(collections),
    llm_count=len(llm_models),
    embedding_count=len(embedding_models),
)
render_status_strip(base_url=st.session_state.base_url, collections_count=len(collections))

boot_report = st.session_state.lmstudio_boot_report
if isinstance(boot_report, dict) and boot_report.get("message"):
    if boot_report.get("ok"):
        st.success("LM Studio: " + str(boot_report.get("message")))
    else:
        st.warning("LM Studio: " + str(boot_report.get("message")))

st.markdown(
    """
    <div class="footer-note">
        The <strong>RAG method does not fine-tune</strong> the model. It creates a local knowledge base from your files
        and uses the LM Studio model to answer questions about them with citations.
    </div>
    """,
    unsafe_allow_html=True,
)

with st.sidebar:
    st.header("⚙️ LM Studio connection")

    st.session_state.base_url = st.text_input("Base URL", value=st.session_state.base_url)
    st.session_state.api_key = st.text_input("API Key", value=st.session_state.api_key, type="password")
    st.session_state.timeout_seconds = st.number_input(
        "Timeout (sec.)",
        min_value=30,
        max_value=900,
        value=int(st.session_state.timeout_seconds),
        step=10,
    )

    start_col, check_col = st.columns(2)
    with start_col:
        if st.button("🚀 Start LM Studio", type="primary", use_container_width=True):
            report = ensure_lmstudio_online(
                st.session_state.base_url,
                api_key=st.session_state.api_key,
                wait_timeout=LMSTUDIO_START_TIMEOUT,
            )
            st.session_state.lmstudio_boot_report = report
            if report.get("ok"):
                try:
                    refresh_models()
                except Exception:
                    pass
                st.success(report.get("message", "LM Studio is ready."))
            else:
                st.warning(report.get("message", "Automatic startup was not possible."))

    with check_col:
        if st.button("🔎 Check / connect", use_container_width=True):
            report = ensure_lmstudio_online(
                st.session_state.base_url,
                api_key=st.session_state.api_key,
                wait_timeout=15,
            )
            st.session_state.lmstudio_boot_report = report
            if report.get("ok"):
                try:
                    refresh_models()
                except Exception:
                    pass
                st.success("LM Studio is responding normally.")
            else:
                st.error(report.get("message", "Failed to connect to LM Studio."))

    if st.button("🔄 Refresh models", use_container_width=True):
        try:
            refresh_models()
            st.success("The model list was refreshed.")
        except Exception as exc:
            st.error(f"Failed to retrieve models: {exc}")

    report = st.session_state.lmstudio_boot_report
    if isinstance(report, dict):
        status_label = "🟢 Service active" if report.get("ok") else "🟠 Service offline / unknown"
        st.markdown(
            f"""
            <div class="service-box">
                <strong>{status_label}</strong><br>
                Host: {report.get("host", "-")} • Port: {report.get("port", "-")}<br>
                {report.get("message", "")}
            </div>
            """,
            unsafe_allow_html=True,
        )
        actions = report.get("actions") or []
        if actions:
            st.caption("Latest startup actions:")
            for action in actions:
                st.write(f"- {action}")

    st.divider()
    st.subheader("🎯 Service model selection")

    cached_llms = get_cached_models("llm")
    cached_embeddings = get_cached_models("embedding")
    service_llm_options = [AUTO_OPTION] + [model.identifier for model in cached_llms]

    if st.session_state.service_llm_model not in service_llm_options:
        st.session_state.service_llm_model = AUTO_OPTION

    st.selectbox(
        "LLM to load into the service",
        options=service_llm_options,
        index=service_llm_options.index(st.session_state.service_llm_model),
        format_func=format_model_option,
        key="service_llm_model",
        help="The list comes from the API when the service is active, or from the lms CLI when the service is offline.",
    )

    service_btn_col1, service_btn_col2 = st.columns(2)
    with service_btn_col1:
        if st.button("🚀 Service + model", type="primary", use_container_width=True):
            try:
                if len(service_llm_options) <= 1:
                    raise LMStudioError("No LLM models were found for selection. Open LM Studio at least once and download a model.")

                report = ensure_lmstudio_online(
                    st.session_state.base_url,
                    api_key=st.session_state.api_key,
                    wait_timeout=LMSTUDIO_START_TIMEOUT,
                )
                st.session_state.lmstudio_boot_report = report
                if not report.get("ok"):
                    raise LMStudioError(report.get("message", "Failed to start the service."))

                refresh_models()
                client = build_client()
                selected_service_model, selected_service_reason = resolve_model(
                    model_type="llm",
                    selected_identifier=st.session_state.service_llm_model,
                    engine=engine,
                )
                load_message = ensure_model_loaded(client, selected_service_model)
                st.session_state.ask_llm_model = selected_service_model.identifier
                st.success(f"The service is active and is using model `{selected_service_model.identifier}`.")
                st.caption(selected_service_reason)
                st.caption(load_message)
                st.rerun()
            except Exception as exc:
                st.error(f"Failed to start service/model: {exc}")

    with service_btn_col2:
        if st.button("📥 Load selected", use_container_width=True):
            try:
                if st.session_state.service_llm_model == AUTO_OPTION:
                    raise LMStudioError("First select a specific LLM from the list, or click 'Service + model' for automatic selection.")
                report = ensure_lmstudio_online(
                    st.session_state.base_url,
                    api_key=st.session_state.api_key,
                    wait_timeout=20,
                )
                st.session_state.lmstudio_boot_report = report
                if not report.get("ok"):
                    raise LMStudioError(report.get("message", "Failed to connect to the service."))

                refresh_models()
                client = build_client()
                chosen_model = get_model_by_identifier(st.session_state.service_llm_model)
                if not chosen_model:
                    raise LMStudioError("The selected model was not found after refreshing the list.")
                load_message = ensure_model_loaded(client, chosen_model)
                st.session_state.ask_llm_model = chosen_model.identifier
                st.success(load_message)
                st.rerun()
            except Exception as exc:
                st.error(f"Failed to load model: {exc}")

    st.divider()
    st.subheader("📦 Available models")
    st.caption(f"Model list source: {st.session_state.models_source}")
    st.write(f"LLM models: **{len(llm_models)}**")
    st.write(f"Embedding models: **{len(embedding_models)}**")

    if cached_llms:
        st.caption("LLM:")
        for model in cached_llms:
            badge = "🟢 loaded" if model.is_loaded else "⚪ available"
            st.write(f"- `{model.identifier}` — {badge}")

    if cached_embeddings:
        st.caption("Embeddings:")
        for model in cached_embeddings:
            badge = "🟢 loaded" if model.is_loaded else "⚪ available"
            st.write(f"- `{model.identifier}` — {badge}")

    if not cached_embeddings:
        st.warning(
            "No embedding model was found in LM Studio. "
            "Download or load an embedding model and click 'Refresh models' again."
        )


tab_overview, tab_build, tab_ask, tab_collections = st.tabs([
    "🏠 Overview",
    "🧱 Build knowledge",
    "❓ Ask",
    "🗂️ Collections",
])

with tab_overview:
    st.subheader("Main dashboard")
    left, right = st.columns([1.05, 1.15])

    with left:
        st.markdown(
            """
            <div class="mini-card">
                <div class="mini-card-title">Main features</div>
                <div style="color:#0f172a; line-height:1.75;">
                    • Local RAG on your files<br>
                    • LLM selection and automatic service startup<br>
                    • Scrollable answer panel with copy/download<br>
                    • Attach files to the question (PDF, DOCX, code, etc.)<br>
                    • Knowledge collections with embeddings and search<br>
                    • Clean display of sources, collections, and attachments
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        st.markdown(
            """
            <div class="mini-card">
                <div class="mini-card-title">Quick usage flow</div>
                <div style="color:#0f172a; line-height:1.75;">
                    1. Check LM Studio from the sidebar.<br>
                    2. Upload files or build a knowledge collection.<br>
                    3. Ask a question and optionally add attachments.<br>
                    4. View or copy the answer from the panel.
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    with right:
        if st.session_state.latest_result:
            latest_question = st.session_state.latest_result.get("question", "-")
            latest_answer = st.session_state.latest_result.get("answer", "")
            st.markdown(
                """
                <div class="mini-card">
                    <div class="mini-card-title">Latest answer</div>
                </div>
                """,
                unsafe_allow_html=True,
            )
            st.caption("Question: " + ellipsize_middle(latest_question, 90))
            st.info((latest_answer[:220] + "…") if len(latest_answer) > 320 else latest_answer)
            st.caption("The full formatted answer appears in the 'Ask' tab.")
        else:
            st.info("There is no answer yet. Ask a question from the 'Ask' tab.")

with tab_build:
    st.subheader("1) Build / update knowledge")

    collection_name = st.text_input("Collection name", value="knowledge_base")
    build_embedding_options = [AUTO_OPTION] + embedding_models
    selected_build_embedding_model = st.selectbox(
        "Embedding model from LM Studio",
        options=build_embedding_options,
        index=0,
        format_func=format_model_option,
        key="build_embedding_model",
        help="Leave it on automatic selection to find and load an embedding model automatically.",
    )

    uploaded_files = st.file_uploader(
        "Knowledge files",
        type=None,
        accept_multiple_files=True,
        help="Supported: txt, md, py, json, csv, html, pdf, docx",
    )

    col_build_1, col_build_2, col_build_3 = st.columns(3)
    with col_build_1:
        chunk_size = st.slider("Chunk size", min_value=300, max_value=2500, value=DEFAULT_CHUNK_SIZE, step=50)
    with col_build_2:
        chunk_overlap = st.slider("Chunk overlap", min_value=0, max_value=500, value=DEFAULT_CHUNK_OVERLAP, step=10)
    with col_build_3:
        batch_size = st.slider("Batch size embeddings", min_value=1, max_value=64, value=DEFAULT_EMBED_BATCH_SIZE, step=1)

    try:
        preview_embedding, preview_reason = resolve_model(
            model_type="embedding",
            selected_identifier=selected_build_embedding_model,
            engine=engine,
        )
        st.caption(f"Automatic resolution: `{preview_embedding.identifier}` — {preview_reason}")
    except Exception as exc:
        st.caption(f"Automatic embedding model resolution: {exc}")

    if uploaded_files:
        valid_count = len([file for file in uploaded_files if is_supported(file.name)])
        invalid_files = [file.name for file in uploaded_files if not is_supported(file.name)]
        c1, c2 = st.columns(2)
        c1.metric("Valid files", valid_count)
        c2.metric("Invalid files", len(invalid_files))
        if invalid_files:
            st.warning("Unsupported files: " + ", ".join(invalid_files))

    if st.button("Build / rebuild RAG", type="primary", use_container_width=True):
        try:
            valid_files = [file for file in (uploaded_files or []) if is_supported(file.name)]
            if not valid_files:
                raise ValueError("You must upload at least one supported file.")

            client = build_client()
            embedding_model, embedding_reason = resolve_model(
                model_type="embedding",
                selected_identifier=selected_build_embedding_model,
                engine=engine,
            )
            load_message = ensure_model_loaded(client, embedding_model)

            with st.spinner("Extracting text, chunking, and creating embeddings..."):
                metadata = engine.build_collection(
                    collection_name=collection_name,
                    uploaded_files=valid_files,
                    client=client,
                    embedding_model=embedding_model.identifier,
                    chunk_size=chunk_size,
                    chunk_overlap=chunk_overlap,
                    batch_size=batch_size,
                )

            st.success("The knowledge base was created successfully.")
            st.info(f"Embedding model: `{embedding_model.identifier}`")
            st.caption(embedding_reason)
            st.caption(load_message)
            st.json(metadata)
            st.rerun()
        except Exception as exc:
            st.error(f"Failed to create collection: {exc}")

with tab_ask:
    st.subheader("2) Ask a question about the knowledge base")

    collection_options = [NO_COLLECTION_OPTION] + collections
    selected_collection_option = st.selectbox(
        "Collection",
        options=collection_options,
        index=0 if not collections else 1,
        format_func=format_collection_option,
        help="You can answer either from an existing collection or only from question attachments.",
    )
    selected_collection = None if selected_collection_option == NO_COLLECTION_OPTION else selected_collection_option

    llm_options = [AUTO_OPTION] + llm_models
    embedding_options = [AUTO_OPTION] + embedding_models

    if st.session_state.ask_llm_model not in llm_options:
        st.session_state.ask_llm_model = AUTO_OPTION

    ask_col1, ask_col2 = st.columns(2)
    with ask_col1:
        selected_llm_model = st.selectbox(
            "LLM model from LM Studio",
            options=llm_options,
            index=llm_options.index(st.session_state.ask_llm_model),
            format_func=format_model_option,
            key="ask_llm_model",
            help="Leave it on automatic selection so it can choose an available LLM by itself.",
        )
    with ask_col2:
        qa_embedding_model = st.selectbox(
            "Embedding model for search",
            options=embedding_options,
            index=0,
            format_func=format_model_option,
            key="qa_embedding_model",
            help="It usually needs to match the collection's embedding model. Automatic selection handles that when possible.",
        )

    ask_col3, ask_col4, ask_col5 = st.columns(3)
    with ask_col3:
        top_k = st.slider("Top-K relevant excerpts", min_value=1, max_value=12, value=DEFAULT_TOP_K, step=1)
    with ask_col4:
        temperature = st.slider("Temperature", min_value=0.0, max_value=1.0, value=DEFAULT_TEMPERATURE, step=0.05)
    with ask_col5:
        max_tokens = st.slider("Max tokens", min_value=128, max_value=4096, value=DEFAULT_MAX_TOKENS, step=64)

    question = st.text_area(
        "Question",
        value="",
        height=180,
        placeholder="Write your question here about your own data...",
    )

    question_attachments = st.file_uploader(
        "Question attachments",
        type=None,
        accept_multiple_files=True,
        key="qa_attachments",
        help="You can attach PDF, DOCX, TXT, JSON, CSV, and many code/text file types. Attachments are inserted directly into the prompt as additional context.",
    )

    if question_attachments:
        valid_question_attachments = [file for file in question_attachments if is_supported(file.name)]
        invalid_question_attachments = [file.name for file in question_attachments if not is_supported(file.name)]
        qa_stat1, qa_stat2 = st.columns(2)
        qa_stat1.metric("Valid attachments", len(valid_question_attachments))
        qa_stat2.metric("Unsupported", len(invalid_question_attachments))
        if valid_question_attachments:
            st.caption("They will be analyzed in the question: " + ", ".join(file.name for file in valid_question_attachments))
        if invalid_question_attachments:
            st.warning("Unsupported attachments: " + ", ".join(invalid_question_attachments))

    try:
        preview_llm, preview_llm_reason = resolve_model(
            model_type="llm",
            selected_identifier=selected_llm_model,
            engine=engine,
            collection_name=selected_collection,
        )
        st.caption(f"LLM: `{preview_llm.identifier}` — {preview_llm_reason}")
    except Exception as exc:
        st.caption(f"Automatic LLM resolution: {exc}")

    if selected_collection:
        try:
            preview_qa_embedding, preview_qa_embedding_reason = resolve_model(
                model_type="embedding",
                selected_identifier=qa_embedding_model,
                engine=engine,
                collection_name=selected_collection,
            )
            st.caption(f"Search embedding: `{preview_qa_embedding.identifier}` — {preview_qa_embedding_reason}")
        except Exception as exc:
            st.caption(f"Automatic embedding resolution: {exc}")
    else:
        st.caption("Search embedding: not required when you work only with question attachments.")

    if st.button("Ask the model", type="primary", use_container_width=True):
        try:
            if not question.strip():
                raise ValueError("The question is empty.")

            attachment_documents, attachment_warnings = load_question_attachments(question_attachments or [])
            if not selected_collection and not attachment_documents:
                raise ValueError("Select a collection or upload at least one readable attachment.")

            client = build_client()
            llm_model, llm_reason = resolve_model(
                model_type="llm",
                selected_identifier=selected_llm_model,
                engine=engine,
                collection_name=selected_collection,
            )

            embedding_model = None
            embedding_reason = "No embedding model was required because the answer was based only on attachments."
            embedding_load_message = ""

            if selected_collection:
                embedding_model, embedding_reason = resolve_model(
                    model_type="embedding",
                    selected_identifier=qa_embedding_model,
                    engine=engine,
                    collection_name=selected_collection,
                )

            llm_load_message = ensure_model_loaded(client, llm_model)
            if embedding_model is not None:
                embedding_load_message = ensure_model_loaded(client, embedding_model)

            with st.spinner("Searching relevant chunks, analyzing attachments, and generating the answer..."):
                result = engine.answer_question(
                    collection_name=selected_collection,
                    question=question,
                    client=client,
                    llm_model=llm_model.identifier,
                    embedding_model=embedding_model.identifier if embedding_model else None,
                    top_k=top_k,
                    temperature=temperature,
                    max_tokens=max_tokens,
                    attachment_documents=attachment_documents,
                )

            result["question"] = question
            result["meta"] = {
                "collection": selected_collection or "-",
                "llm_model": llm_model.identifier,
                "embedding_model": embedding_model.identifier if embedding_model else "-",
                "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "attachments": [doc.filename for doc in attachment_documents],
                "attachments_count": len(attachment_documents),
            }
            st.session_state.latest_result_seq = int(st.session_state.latest_result_seq) + 1
            result["result_seq"] = st.session_state.latest_result_seq
            st.session_state.latest_result = result

            st.success("The answer was generated.")
            st.caption(f"LLM: `{llm_model.identifier}` — {llm_reason}")
            st.caption(llm_load_message)
            if embedding_model is not None:
                st.caption(f"Embedding: `{embedding_model.identifier}` — {embedding_reason}")
                st.caption(embedding_load_message)
            else:
                st.caption(f"Embedding: {embedding_reason}")
            for warning in attachment_warnings:
                st.warning(warning)
            st.rerun()
        except Exception as exc:
            st.error(f"Failed to generate answer: {exc}")

    st.markdown("---")
    render_answer_panel(st.session_state.latest_result, panel_key_prefix="ask")

with tab_collections:
    st.subheader("3) Existing collections")

    if not collections:
        st.warning("There are no knowledge collections yet.")
    else:
        search_filter = st.text_input("Collection filter", value="")
        visible_collections = [
            collection for collection in collections
            if search_filter.strip().lower() in collection.lower()
        ]

        if not visible_collections:
            st.info("No collections were found with this filter.")

        for collection in visible_collections:
            try:
                metadata = engine.get_collection_metadata(collection)
                with st.container(border=True):
                    c1, c2 = st.columns([2.4, 1])
                    with c1:
                        st.markdown(f"### {ellipsize_middle(collection, 42)}")
                        st.caption(f"Full name: {collection}")
                        st.write(f"Embedding model: `{ellipsize_middle(metadata.get('embedding_model', '-'), 52)}`")
                        st.write(f"Files: {len(metadata.get('documents', []))}")
                        st.write(f"Chunks: {metadata.get('num_chunks', 0)}")
                        st.write(f"Created: {metadata.get('created_at', '-')}")
                        if metadata.get("documents"):
                            shown_docs = [ellipsize_middle(doc, 54) for doc in metadata["documents"]]
                            st.caption("• " + "\n• ".join(shown_docs))
                    with c2:
                        if st.button("🗑️ Delete", key=f"delete_{collection}", use_container_width=True, help=collection):
                            engine.delete_collection(collection)
                            st.success(f"Collection '{collection}' was deleted.")
                            st.rerun()
            except Exception as exc:
                st.error(f"Error in collection '{collection}': {exc}")

st.divider()
st.markdown(
    """
    <div class="footer-note">
        <strong>Note:</strong> If you want real training / fine-tuning (e.g. LoRA or QLoRA),
        that is a different pipeline from RAG. This application helps you make
        the LM Studio model answer questions about your own files without changing the model weights.
    </div>
    """,
    unsafe_allow_html=True,
)